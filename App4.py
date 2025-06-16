import pandas as pd
import streamlit as st
from docx import Document

# Fonction pour générer un rapport Word par CAPA
def generate_report(df):
    for capa in df['Numéro CAPA'].unique():
        capa_data = df[df['Numéro CAPA'] == capa]
        doc = Document()
        doc.add_heading(f'CAPA Report: {capa}', 0)

        for index, row in capa_data.iterrows():
            doc.add_heading(f"Détails pour la CAPA : {row['Numéro CAPA']}", level=1)
            for col in df.columns:
                doc.add_paragraph(f"{col} : {row[col]}")
        
        doc.save(f"CAPA_Report_{capa}.docx")

# Interface Streamlit
st.title('📝 Générateur de rapports CAPA')

uploaded_file = st.file_uploader("📂 Charger un fichier Excel", type="xlsx")

if uploaded_file:
    df = pd.read_excel(uploaded_file, engine='openpyxl')
    st.dataframe(df)

    # Filtres dynamiques
    status = st.multiselect('Statut', df['Statut'].dropna().unique())
    anomaly_type = st.multiselect("Type d'anomie", df["Type d'anomie"].dropna().unique())
    remontee = st.multiselect('Remontée', df['Remontée'].dropna().unique())
    capa_number = st.multiselect('Numéro CAPA', df['Numéro CAPA'].dropna().unique())
    instrument = st.multiselect('Instrument', df['Instrument'].dropna().unique())
    ipr = st.multiselect('IPR', df['IPR'].dropna().unique())

    # Application des filtres
    if status:
        df = df[df['Statut'].isin(status)]
    if anomaly_type:
        df = df[df["Type d'anomie"].isin(anomaly_type)]
    if remontee:
        df = df[df['Remontée'].isin(remontee)]
    if capa_number:
        df = df[df['Numéro CAPA'].isin(capa_number)]
    if instrument:
        df = df[df['Instrument'].isin(instrument)]
    if ipr:
        df = df[df['IPR'].isin(ipr)]

    st.dataframe(df)

    if st.button('📄 Générer les rapports Word'):
        generate_report(df)
        st.success('✅ Rapports générés avec succès !')
